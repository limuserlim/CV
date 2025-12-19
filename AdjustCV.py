import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import PIL.Image
import os

# הגדרות דף
st.set_page_config(page_title="מחולל קורות חיים AI", layout="centered")
st.title("📄 התאמת קורות חיים אישית")

# --- הגדרות API ---
API_KEY = API_KEY = st.secrets["GOOGLE_API_KEY"]
genai.configure(api_key=API_KEY)

def run_ai_logic(cv_text, job_input):
    model = genai.GenerativeModel("models/gemini-2.0-flash-exp")
    prompt = f"התאם את קורות החיים למודעה. השתמש רק במידע הקיים. עברית בלבד.\nCV: {cv_text}\nJob: {job_input}"
    
    # בדיקה אם מדובר בתמונה (Streamlit מעלה קבצים כ-Bytes)
    if hasattr(job_input, 'type') and job_input.type.startswith('image'):
        img = PIL.Image.open(job_input)
        response = model.generate_content([prompt, img])
    else:
        response = model.generate_content(prompt)
    
    return response.text

def create_docx(text):
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 == 1: run.bold = True
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- ממשק המשתמש ---
if not os.path.exists("CVall.txt"):
    st.error("שגיאה: הקובץ CVall.txt לא נמצא בתיקייה!")
else:
    with open("CVall.txt", "r", encoding="utf-8") as f:
        cv_content = f.read()

    st.info("קובץ המקור CVall.txt נטען בהצלחה.")

    # בחירת סוג הקלט
    input_type = st.radio("איך תרצי להזין את המודעה?", ["טקסט/קישור", "תמונה (צילום מסך)"])

    job_data = None
    if input_type == "טקסט/קישור":
        job_data = st.text_area("הדביקי כאן את המודעה:")
    else:
        job_data = st.file_uploader("העלי צילום מסך של המודעה", type=['png', 'jpg', 'jpeg'])

    if st.button("בצע התאמה וצור קובץ Word"):
        if job_data:
            with st.spinner("ג'מיני מעבד את הנתונים..."):
                try:
                    result_text = run_ai_logic(cv_content, job_data)
                    st.subheader("התוצאה:")
                    st.markdown(result_text)
                    
                    # יצירת הורדה
                    docx_file = create_docx(result_text)
                    st.download_button(
                        label="⬇️ הורד קורות חיים בפורמט Word",
                        data=docx_file,
                        file_name="CV_Adjusted.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"קרתה שגיאה: {e}")
        else:
            st.warning("בבקשה הזיני מודעה קודם.")
